import streamlit as st
import pandas as pd
import PyPDF2
from docx import Document
import os
from google.cloud import language_v1

# Set up Google API Key
google_api_key = os.getenv("GOOGLE_API_KEY")
if not google_api_key:
    st.error("Google API Key is not set. Please configure it in your environment variables.")

client = language_v1.LanguageServiceClient()

def extract_text_from_pdf(pdf_file):
    reader = PyPDF2.PdfReader(pdf_file)
    text = "".join([page.extract_text() for page in reader.pages if page.extract_text()])
    return text

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def analyze_text_with_google(text):
    document = language_v1.Document(content=text, type_=language_v1.Document.Type.PLAIN_TEXT)
    response = client.analyze_entities(request={"document": document})
    keywords = [entity.name for entity in response.entities]
    return keywords

def process_price_list(uploaded_files):
    data_frames = []
    for file in uploaded_files:
        if file.name.endswith('.csv'):
            df = pd.read_csv(file)
        elif file.name.endswith('.xlsx'):
            df = pd.read_excel(file)
        else:
            st.error(f"Unsupported file format: {file.name}")
            continue
        data_frames.append(df)
    if data_frames:
        combined_df = pd.concat(data_frames, ignore_index=True)
        st.write("### Combined Price List")
        st.write(combined_df)
        
        st.write("### Price Analysis")
        st.write(combined_df.describe())

def process_research_paper(uploaded_file):
    if uploaded_file.name.endswith('.pdf'):
        text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.name.endswith('.docx'):
        text = extract_text_from_docx(uploaded_file)
    else:
        st.error("Unsupported file format")
        return
    summary = " ".join(analyze_text_with_google(text)[:10])
    st.write("### Key Topics & Summary")
    st.write(summary)

def process_resumes(uploaded_files, job_description):
    matched_resumes = []
    job_keywords = set(analyze_text_with_google(job_description))
    
    for file in uploaded_files:
        if file.name.endswith('.pdf'):
            text = extract_text_from_pdf(file)
        elif file.name.endswith('.docx'):
            text = extract_text_from_docx(file)
        else:
            st.error(f"Unsupported file format: {file.name}")
            continue
        resume_keywords = set(analyze_text_with_google(text))
        similarity = len(job_keywords & resume_keywords) / max(len(job_keywords), 1)
        matched_resumes.append((file.name, similarity))
    
    matched_resumes.sort(key=lambda x: x[1], reverse=True)
    st.write("### Matched Resumes")
    for name, score in matched_resumes:
        st.write(f"{name}: {score:.2f}")

def main():
    st.title("Advanced Document Analyzer")
    option = st.sidebar.selectbox("Choose a scenario", ["Price List Analyzer", "Research Paper Simplifier", "Resume Matcher for Hiring"])
    
    if option == "Price List Analyzer":
        uploaded_files = st.file_uploader("Upload price lists (CSV or Excel)", type=["csv", "xlsx"], accept_multiple_files=True)
        if uploaded_files:
            process_price_list(uploaded_files)
    
    elif option == "Research Paper Simplifier":
        uploaded_file = st.file_uploader("Upload research paper (PDF or DOCX)", type=["pdf", "docx"])
        if uploaded_file:
            process_research_paper(uploaded_file)
    
    elif option == "Resume Matcher for Hiring":
        job_description = st.text_area("Enter job description")
        uploaded_files = st.file_uploader("Upload resumes (PDF or DOCX)", type=["pdf", "docx"], accept_multiple_files=True)
        if uploaded_files and job_description:
            process_resumes(uploaded_files, job_description)

if __name__ == "__main__":
    main()
